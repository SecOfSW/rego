from docx import Document
from docx.shared import Pt, RGBColor
import winreg
import REGO_reg
import data

# JSON file format
"""
[
    {
        "HIVE": string,
        "PATH": string,
        "TYPE": string,
        "VULN_VAL": list,
        "SAFE_VAL": list
    }
]
"""

MAX_NUM = 100000
class REGO_reg_scan(REGO_reg.REGO_reg):
    def __init__(self):
        pass
    
    def scan(self):
        self.init_docx()

        # Information of registries with vulnerable value
        scan_result = []
        for i, m in zip(range(MAX_NUM),data.sec_reg.reg_info_list):
            NAME = m.NAME
            HIVE = m.HIVE
            PATH = m.PATH
            ATTRIBUTE = m.ATTRIBUTE
            VULN_VAL = m.VULN_VAL
            SAFE_VAL = m.SAFE_VAL
            DESC = m.DESC
            
            attList = []
            try:
                attList =  self.getReg(HIVE=HIVE, PATH=PATH)
            except Exception as e:
                print("[{0}] {1:130} - NOT found".format(str(i), REGO_reg.hiveIntToStr(HIVE)+"\\"+str(PATH)+"\\"+str(ATTRIBUTE)))
                continue

            for _ in attList:
                if _[0] == ATTRIBUTE:
                    print("[{0}] {1:130} - <CURRENT_VAL>: {2:>10}, <SAFE_VAL>: {3:>10}, <VULN_VAL>: {4:>10}, <DESC>: {5}".format(str(i), REGO_reg.hiveIntToStr(HIVE)+"\\"+str(PATH)+"\\"+str(ATTRIBUTE), _[1], str(SAFE_VAL), str(VULN_VAL), DESC)) 
                    self.input_docx(NAME, REGO_reg.hiveIntToStr(HIVE), PATH, ATTRIBUTE, _[1], SAFE_VAL, DESC)
                    if _[1] not in SAFE_VAL:
                        scan_result.append({
                                            "HIVE": REGO_reg.hiveIntToStr(HIVE), 
                                            "PATH": PATH, 
                                            "ATTRIBUTE": ATTRIBUTE,
                                            "CURRENT_VALUE": _[1],
                                            "SAFE_VALUE": SAFE_VAL
                                            })
                    break
        self.end_docx()
        return scan_result
        
    def init_docx(self):
        self.document = Document()
        self.document.styles['Heading 1'].font.size = Pt(18)
        self.document.styles['Heading 2'].font.size = Pt(14)
        self.document.add_heading("REGO : Scan Report", 0)
        p = self.document.add_paragraph('Result of REGO_Scan.\n')
        p.add_run('Security related registery checked.')
        self.document.add_heading('Checked Registery', level=1)

    def end_docx(self):
        self.document.add_page_break()
        self.document.save('report.docx')

    def input_docx(self, NAME, HIVE, PATH, ATTRIBUTE, CURRENT_VALUE, SAFE_VALUE, DESC):
        h = self.document.add_heading(NAME, level=2)
        if CURRENT_VALUE in SAFE_VALUE:
            h.add_run(" SAFE").font.color.rgb = RGBColor(0x22, 0x8B, 0x22)
        else:
            h.add_run(" DANGER").font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        p = self.document.add_paragraph("[HIVE]\n" + str(HIVE) + '\n\n')
        p.add_run("[PATH]\n" + str(PATH) + '\n\n')
        p.add_run("[ATTRIBUTE]\n" + str(ATTRIBUTE) + '\n\n')
        p.add_run("[CURRENT_VALUE]\n" + str(CURRENT_VALUE) + '\n\n')
        p.add_run("[SAFE_VALUE]\n" + str(SAFE_VALUE) + '\n\n')
        p.add_run("[DESC]\n" + str(DESC) + '\n\n')
        return